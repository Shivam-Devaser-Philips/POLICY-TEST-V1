from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

old_text = [
    ("1.1 Purpose", "This policy provides guidance for employee leave and code of conduct."),
    ("1.2 Scope", "Applies to all employees in the organization worldwide."),
    ("2.1 Annual Leave", "Employees are entitled to 20 days of annual leave each year."),
    ("2.2 Sick Leave", "Employees may use up to 10 days of paid sick leave. "),
    ("3.1 Data Privacy", "The company collects only necessary employee data and protects privacy."),
]

new_text = [
    ("1.1 Purpose", "This policy provides guidance for employee leave, compliance, and modern workplace practices."),
    ("1.2 Scope", "Applies to employees and contractors in all regions."),
    ("2.1 Annual Leave", "Employees are entitled to 22 days of annual leave each year."),
    ("2.2 Sick Leave", "Employees may use up to 12 days of paid sick leave."),
    ("2.3 Compliance", "Leave policies must follow labor regulations and audit requirements."),
    ("3.1 Data Privacy", "The company collects only necessary employee data and ensures privacy and security."),
]


def write_docx(filename, records):
    doc = Document()
    doc.add_heading('Policy Document', level=1)
    for heading, body in records:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)
    doc.save(filename)


def write_pdf(filename, records):
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    y = height - 72
    c.setFont('Helvetica-Bold', 16)
    c.drawString(72, y, 'Policy Document')
    y -= 36
    c.setFont('Helvetica', 11)
    for heading, body in records:
        if y < 108:
            c.showPage()
            y = height - 72
            c.setFont('Helvetica', 11)
        c.setFont('Helvetica-Bold', 12)
        c.drawString(72, y, heading)
        y -= 18
        c.setFont('Helvetica', 10)
        for line in body.split('. '):
            if not line.strip():
                continue
            wrapped = line.strip() + '.'
            c.drawString(80, y, wrapped)
            y -= 14
            if y < 72:
                c.showPage()
                y = height - 72
                c.setFont('Helvetica', 10)
        y -= 10
    c.save()


if __name__ == '__main__':
    write_docx('data/sample_old.docx', old_text)
    write_docx('data/sample_new.docx', new_text)
    write_pdf('data/sample_old.pdf', old_text)
    write_pdf('data/sample_new.pdf', new_text)
    print('Sample documents created in data/')
